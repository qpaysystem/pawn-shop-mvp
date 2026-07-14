#!/usr/bin/env python3
"""Generate Lombard post-implementation Word document from markdown."""

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_formatted_runs(paragraph, text: str, base_bold: bool = False) -> None:
  """Add text with **bold** and *italic* markers."""
  pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|«[^»]+»)")
  pos = 0
  for match in pattern.finditer(text):
    if match.start() > pos:
      run = paragraph.add_run(text[pos:match.start()])
      run.bold = base_bold
    chunk = match.group()
    if chunk.startswith("**"):
      run = paragraph.add_run(chunk[2:-2])
      run.bold = True
    elif chunk.startswith("*"):
      run = paragraph.add_run(chunk[1:-1])
      run.italic = True
    else:
      run = paragraph.add_run(chunk)
      run.bold = base_bold
    pos = match.end()
  if pos < len(text):
    run = paragraph.add_run(text[pos:])
    run.bold = base_bold


def parse_table_row(line: str) -> list[str]:
  return [c.strip() for c in line.strip().strip("|").split("|")]


def is_separator_row(cells: list[str]) -> bool:
  return all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) or c == "" for c in cells)


def add_table(doc: Document, rows: list[list[str]]) -> None:
  if not rows:
    return
  col_count = max(len(r) for r in rows)
  table = doc.add_table(rows=len(rows), cols=col_count)
  table.style = "Table Grid"
  table.autofit = True

  for r_idx, row in enumerate(rows):
    for c_idx in range(col_count):
      cell = table.rows[r_idx].cells[c_idx]
      value = row[c_idx] if c_idx < len(row) else ""
      p = cell.paragraphs[0]
      p.clear()
      add_formatted_runs(p, value, base_bold=(r_idx == 0))
      for run in p.runs:
        run.font.size = Pt(10)
      if r_idx == 0:
        set_cell_shading(cell, "E8EEF7")


def build_doc(md_path: Path, out_path: Path) -> None:
  text = md_path.read_text(encoding="utf-8")
  lines = text.splitlines()

  doc = Document()
  section = doc.sections[0]
  section.top_margin = Cm(2)
  section.bottom_margin = Cm(2)
  section.left_margin = Cm(2.5)
  section.right_margin = Cm(2)

  normal = doc.styles["Normal"]
  normal.font.name = "Calibri"
  normal.font.size = Pt(11)
  normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

  i = 0
  table_buffer: list[list[str]] = []
  in_code_block = False
  code_lines: list[str] = []

  def flush_code_block() -> None:
    nonlocal code_lines
    if not code_lines:
      return
    for code_line in code_lines:
      p = doc.add_paragraph()
      p.paragraph_format.left_indent = Cm(0.8)
      p.paragraph_format.space_after = Pt(0)
      run = p.add_run(code_line)
      run.font.name = "Consolas"
      run.font.size = Pt(9.5)
      run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    code_lines = []
    doc.add_paragraph()

  def flush_table() -> None:
    nonlocal table_buffer
    if table_buffer:
      add_table(doc, table_buffer)
      table_buffer = []
      doc.add_paragraph()

  while i < len(lines):
    line = lines[i]
    stripped = line.strip()

    if stripped.startswith("```"):
      if in_code_block:
        flush_code_block()
        in_code_block = False
      else:
        flush_table()
        in_code_block = True
      i += 1
      continue

    if in_code_block:
      code_lines.append(line.rstrip())
      i += 1
      continue

    if not stripped:
      flush_table()
      i += 1
      continue

    if stripped == "---":
      flush_table()
      i += 1
      continue

    if stripped.startswith("|"):
      cells = parse_table_row(stripped)
      if not is_separator_row(cells):
        table_buffer.append(cells)
      i += 1
      continue

    flush_table()

    if stripped.startswith("# "):
      p = doc.add_heading(stripped[2:], level=0)
      p.alignment = WD_ALIGN_PARAGRAPH.CENTER
      for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
      i += 1
      continue

    if stripped.startswith("## "):
      doc.add_heading(stripped[3:], level=1)
      i += 1
      continue

    if stripped.startswith("### "):
      doc.add_heading(stripped[4:], level=2)
      i += 1
      continue

    if re.match(r"^\d+\.\s", stripped):
      p = doc.add_paragraph(style="List Number")
      content = re.sub(r"^\d+\.\s+", "", stripped)
      add_formatted_runs(p, content)
      i += 1
      continue

    if stripped.startswith("- "):
      p = doc.add_paragraph(style="List Bullet")
      add_formatted_runs(p, stripped[2:])
      i += 1
      continue

    p = doc.add_paragraph()
    add_formatted_runs(p, stripped)
    i += 1

  flush_table()
  flush_code_block()

  doc.core_properties.title = "Как будет работать Lombard после внедрения системы"
  doc.core_properties.author = "Lombard Portal"
  doc.core_properties.comments = "Документ подготовлен до раздела «Поэтапность»"

  doc.save(out_path)


if __name__ == "__main__":
  base = Path(__file__).resolve().parent
  md_src = Path("/tmp/lombard_doc.md")
  out = base / "Lombard-portal-posle-vnedreniya.docx"
  build_doc(md_src, out)
  print(f"Created: {out}")
